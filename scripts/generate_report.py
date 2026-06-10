#!/usr/bin/env python3
import argparse
import json
import platform
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    HRFlowable,
    KeepTogether,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_CONFIG = ROOT / "config" / "weekly_report.json"
DEFAULT_OUTPUT_DIR = ROOT / "outputs"

NAVY = colors.HexColor("#0B2545")
BLUE = colors.HexColor("#2F80ED")
SKY = colors.HexColor("#EAF3FF")
PALE = colors.HexColor("#F6F9FC")
BORDER = colors.HexColor("#D8E2EE")
TEXT = colors.HexColor("#263238")
MUTED = colors.HexColor("#6B7280")


def load_config(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def slug_date(text):
    return text.replace("/", "-").replace(" ", "")


def output_basename(config):
    company = config.get("company", "Agentech")
    reporter = config.get("reporter", "Wesley")
    start = slug_date(config["period"]["start"])
    end = slug_date(config["period"]["end"])
    return f"{company}_周工作汇报_{reporter}_{start}_to_{end}"


def find_existing(paths):
    for path in paths:
        if Path(path).exists():
            return str(path)
    return None


def register_pdf_fonts():
    system = platform.system()
    if system == "Darwin":
        body_path = find_existing(
            [
                "/System/Library/Fonts/Supplemental/Songti.ttc",
                "/System/Library/Fonts/STHeiti Medium.ttc",
            ]
        )
        head_path = find_existing(
            [
                "/System/Library/Fonts/STHeiti Medium.ttc",
                "/System/Library/Fonts/Supplemental/Songti.ttc",
            ]
        )
        body_subfont = 4 if body_path and body_path.endswith("Songti.ttc") else 0
        head_subfont = 0
    elif system == "Windows":
        body_path = find_existing(
            [
                "C:/Windows/Fonts/msyh.ttc",
                "C:/Windows/Fonts/simhei.ttf",
                "C:/Windows/Fonts/simsun.ttc",
            ]
        )
        head_path = body_path
        body_subfont = 0
        head_subfont = 0
    else:
        body_path = find_existing(
            [
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            ]
        )
        head_path = body_path
        body_subfont = 0
        head_subfont = 0

    if not body_path:
        raise RuntimeError("No Chinese font found. Install Noto Sans CJK or run on macOS/Windows.")

    pdfmetrics.registerFont(TTFont("ReportBodyCN", body_path, subfontIndex=body_subfont))
    pdfmetrics.registerFont(TTFont("ReportHeadCN", head_path or body_path, subfontIndex=head_subfont))
    return "ReportBodyCN", "ReportHeadCN"


def para(text, style):
    return Paragraph(str(text).replace("\n", "<br/>"), style)


def build_pdf(config, output_path):
    body_font, head_font = register_pdf_fonts()

    title_style = ParagraphStyle(
        "Title",
        fontName=head_font,
        fontSize=23,
        leading=28,
        alignment=TA_CENTER,
        textColor=colors.white,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        fontName=body_font,
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#DDEBFF"),
    )
    h1_style = ParagraphStyle(
        "H1",
        fontName=head_font,
        fontSize=14,
        leading=18,
        textColor=NAVY,
        spaceBefore=10,
        spaceAfter=5,
    )
    h2_style = ParagraphStyle(
        "H2",
        fontName=head_font,
        fontSize=11.2,
        leading=14.5,
        textColor=NAVY,
    )
    body_style = ParagraphStyle(
        "Body",
        fontName=body_font,
        fontSize=9.2,
        leading=12.2,
        textColor=TEXT,
        alignment=TA_LEFT,
    )
    label_style = ParagraphStyle(
        "Label",
        fontName=head_font,
        fontSize=9,
        leading=12,
        textColor=BLUE,
    )
    field_style = ParagraphStyle(
        "Field",
        fontName=body_font,
        fontSize=9,
        leading=13,
        textColor=TEXT,
    )

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(BORDER)
        canvas.setLineWidth(0.6)
        canvas.line(0.75 * inch, 0.62 * inch, 7.75 * inch, 0.62 * inch)
        canvas.setFont(body_font, 8)
        canvas.setFillColor(MUTED)
        canvas.drawRightString(
            7.75 * inch,
            0.43 * inch,
            f"{config.get('company', 'Agentech')} 周工作汇报 | {config.get('reporter', '')} | 第 {doc.page} 页",
        )
        canvas.restoreState()

    doc = BaseDocTemplate(
        str(output_path),
        pagesize=letter,
        leftMargin=0.72 * inch,
        rightMargin=0.72 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.78 * inch,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=footer)])

    title = f"{config.get('company', 'Agentech')} {config.get('title', '周工作汇报')}"
    period = config.get("period", {})
    subtitle = f"汇报人：{config.get('reporter', '')}　　汇报周期：{period.get('start', '')} - {period.get('end', '')}"
    hero = Table([[para(title, title_style)], [para(subtitle, subtitle_style)]], colWidths=[doc.width])
    hero.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), NAVY),
                ("BOX", (0, 0), (-1, -1), 0, NAVY),
                ("TOPPADDING", (0, 0), (-1, 0), 19),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
                ("TOPPADDING", (0, 1), (-1, 1), 0),
                ("BOTTOMPADDING", (0, 1), (-1, 1), 18),
            ]
        )
    )

    streams = config["workstreams"]
    summary = [[para("序号", label_style), para("工作模块", label_style), para("当前进展", label_style), para("完成度", label_style)]]
    for idx, item in enumerate(streams, start=1):
        summary.append(
            [
                str(idx),
                para(item["name"], body_style),
                para(item.get("current_progress", "待填写"), body_style),
                para(item.get("completion", "____%"), body_style),
            ]
        )
    summary_table = Table(summary, colWidths=[0.48 * inch, 2.95 * inch, 2.1 * inch, 1.0 * inch])
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), SKY),
                ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (0, -1), "CENTER"),
                ("ALIGN", (3, 0), (3, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )

    def section_panel(index, item):
        bullet_text = "<br/>".join([f"• {task}" for task in item.get("tasks", [])])
        progress_lines = []
        if item.get("current_count"):
            progress_lines.append(para(f"当前完成数量<br/>{item['current_count']}", field_style))
        progress_lines.extend(
            [
                para("<b>当前进展</b>", label_style),
                para(item.get("current_progress", "待填写"), field_style),
            ]
        )
        if item.get("notes"):
            progress_lines.append(para(f"<b>补充说明</b><br/>{item['notes']}", field_style))
        progress_lines.append(para(f"完成度<br/>{item.get('completion', '____%')}", field_style))
        progress_box = Table([[progress_lines]], colWidths=[1.68 * inch])
        progress_box.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FBFCFE")),
                    ("BOX", (0, 0), (-1, -1), 0.5, BORDER),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        table = Table(
            [
                [para(f"{index}. {item['name']}", h2_style), ""],
                [para("工作内容", label_style), para("填写区", label_style)],
                [para(bullet_text, body_style), progress_box],
            ],
            colWidths=[4.58 * inch, 1.82 * inch],
        )
        table.setStyle(
            TableStyle(
                [
                    ("SPAN", (0, 0), (1, 0)),
                    ("BACKGROUND", (0, 0), (1, 0), SKY),
                    ("LINEBEFORE", (0, 0), (0, -1), 3, BLUE),
                    ("BOX", (0, 0), (-1, -1), 0.7, BORDER),
                    ("INNERGRID", (0, 1), (-1, -1), 0.35, BORDER),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 9),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                    ("BACKGROUND", (0, 1), (-1, 1), PALE),
                ]
            )
        )
        return KeepTogether([table, Spacer(1, 8)])

    story = [
        hero,
        Spacer(1, 12),
        para("本周工作概览", h1_style),
        HRFlowable(width="100%", thickness=1.0, color=BORDER, spaceBefore=0, spaceAfter=7),
        summary_table,
        Spacer(1, 10),
        para("重点工作进展", h1_style),
        HRFlowable(width="100%", thickness=1.0, color=BORDER, spaceBefore=0, spaceAfter=7),
    ]
    for idx, item in enumerate(streams, start=1):
        story.append(section_panel(idx, item))

    next_table = Table(
        [
            [para("下周计划", label_style), para(config.get("next_week_plan", "待填写"), field_style)],
            [para("需协调事项", label_style), para(config.get("coordination_needed", "待填写"), field_style)],
        ],
        colWidths=[1.15 * inch, 5.35 * inch],
    )
    next_table.setStyle(
        TableStyle(
            [
                ("BOX", (0, 0), (-1, -1), 0.7, BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, BORDER),
                ("BACKGROUND", (0, 0), (0, -1), SKY),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 9),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 9),
            ]
        )
    )
    story.extend(
        [
            Spacer(1, 3),
            para("下周计划与需协调事项", h1_style),
            HRFlowable(width="100%", thickness=1.0, color=BORDER, spaceBefore=0, spaceAfter=7),
            next_table,
        ]
    )
    doc.build(story)


def set_docx_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    for key, value in (
        ("w:ascii", "Calibri"),
        ("w:hAnsi", "Calibri"),
        ("w:eastAsia", "Microsoft YaHei"),
        ("w:cs", "Calibri"),
    ):
        run._element.rPr.rFonts.set(qn(key), value)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def docx_para(doc, text="", size=11, bold=False, color="222222", align=None, style=None):
    paragraph = doc.add_paragraph(style=style)
    if align:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_docx_font(run, size=size, bold=bold, color=color)
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def build_docx(config, output_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    title = f"{config.get('company', 'Agentech')} {config.get('title', '周工作汇报')}"
    period = config.get("period", {})
    docx_para(doc, title, size=22, bold=True, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
    docx_para(
        doc,
        f"汇报人：{config.get('reporter', '')}    汇报周期：{period.get('start', '')} - {period.get('end', '')}",
        size=11,
        color="555555",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    docx_para(doc, "一、本周工作概览", size=16, bold=True, color="2E74B5")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["序号", "工作模块", "当前进展", "完成度"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, "EAF3FF")
        run = cell.paragraphs[0].add_run(text)
        set_docx_font(run, size=10.5, bold=True, color="0B2545")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for idx, item in enumerate(config["workstreams"], start=1):
        row = table.add_row().cells
        values = [str(idx), item["name"], item.get("current_progress", "待填写"), item.get("completion", "____%")]
        for i, value in enumerate(values):
            row[i].text = ""
            run = row[i].paragraphs[0].add_run(value)
            set_docx_font(run, size=10.3)

    docx_para(doc, "二、重点工作进展", size=16, bold=True, color="2E74B5")
    for idx, item in enumerate(config["workstreams"], start=1):
        docx_para(doc, f"{idx}. {item['name']}", size=13, bold=True, color="0B2545")
        docx_para(doc, "工作内容", size=11, bold=True, color="2F80ED")
        for task in item.get("tasks", []):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(task)
            set_docx_font(run, size=10.5)
        if item.get("current_count"):
            docx_para(doc, f"当前完成数量：{item['current_count']}", size=10.5, bold=False)
        docx_para(doc, f"当前进展：{item.get('current_progress', '待填写')}", size=10.5)
        if item.get("notes"):
            docx_para(doc, f"补充说明：{item['notes']}", size=10.5)
        docx_para(doc, f"完成度：{item.get('completion', '____%')}", size=10.5)

    docx_para(doc, "三、下周计划与需协调事项", size=16, bold=True, color="2E74B5")
    docx_para(doc, f"下周计划：{config.get('next_week_plan', '待填写')}", size=10.5)
    docx_para(doc, f"需协调事项：{config.get('coordination_needed', '待填写')}", size=10.5)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="Generate Agentech weekly report PDF and DOCX.")
    parser.add_argument("--config", default=str(DEFAULT_CONFIG), help="Path to weekly_report.json")
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR), help="Output directory")
    args = parser.parse_args()

    config = load_config(args.config)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    base = output_basename(config)
    pdf_path = output_dir / f"{base}.pdf"
    docx_path = output_dir / f"{base}.docx"

    build_pdf(config, pdf_path)
    build_docx(config, docx_path)
    print(f"PDF:  {pdf_path}")
    print(f"DOCX: {docx_path}")


if __name__ == "__main__":
    main()
